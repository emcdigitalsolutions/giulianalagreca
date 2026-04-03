"""
Genera il documento DOCX di presentazione e preventivo
per il sito web Giuliana & La Greca - Onoranze Funebri
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Stili globali ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Margini
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return h

def add_bold_paragraph(bold_text, normal_text=""):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

# ============================================================
# COPERTINA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(80)
run = p.add_run("PROPOSTA COMMERCIALE")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Realizzazione Sito Web Professionale")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run("per")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Giuliana & La Greca")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Agenzia Onoranze Funebri")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Campobello di Licata (AG)")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(60)
run = p.add_run("A cura di: Enrico Maria Caruso")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Full Stack Developer & Web Designer Freelance")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("www.emcdigitalsolutions.it")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

doc.add_page_break()

# ============================================================
# 1. INTRODUZIONE
# ============================================================
add_heading_styled("1. Introduzione", level=1)

doc.add_paragraph(
    "Gentili Titolari dell'Agenzia Onoranze Funebri Giuliana & La Greca,"
)
doc.add_paragraph(
    "con la presente proposta desidero illustrarvi il lavoro di progettazione e sviluppo "
    "del sito web professionale realizzato per la vostra agenzia. Il sito rappresenta "
    "uno strumento fondamentale per la vostra presenza online, offrendo ai clienti un "
    "punto di riferimento digitale accessibile 24 ore su 24, coerente con i valori di "
    "serietà, rispetto e professionalità che contraddistinguono la vostra attività."
)
doc.add_paragraph(
    "Il sito è stato progettato su misura, senza l'utilizzo di template predefiniti "
    "o piattaforme CMS come WordPress, garantendo così massime prestazioni, "
    "leggerezza, velocità di caricamento e totale personalizzazione."
)

# ============================================================
# 2. CARATTERISTICHE DEL SITO
# ============================================================
add_heading_styled("2. Caratteristiche del Sito Web", level=1)

# 2.1
add_heading_styled("2.1 Design Elegante e Su Misura", level=2)
doc.add_paragraph(
    "Il sito è stato sviluppato con un design esclusivo a tema scuro con accenti dorati, "
    "pensato per trasmettere eleganza, sobrietà e rispetto. Ogni elemento visivo è stato "
    "curato nei minimi dettagli:"
)
add_bullet("Palette cromatica raffinata: ", "nero caldo, oro elegante e bianco caldo")
add_bullet("Tipografia professionale: ", "Playfair Display (titoli serif) + Inter (testo sans-serif)")
add_bullet("Animazioni fluide: ", "particelle fluttuanti dorate nella Hero Section")
add_bullet("Effetti di scorrimento: ", "elementi che appaiono gradualmente durante la navigazione")
add_bullet("Banner personalizzato ", "con il logo e il nome dell'agenzia in evidenza")

# 2.2
add_heading_styled("2.2 Struttura e Sezioni", level=2)
doc.add_paragraph("Il sito è strutturato come single-page con navigazione fluida tra le seguenti sezioni:")

sections_data = [
    ("Home / Hero Section", "Presentazione dell'agenzia con banner professionale, logo, "
     "particelle decorative animate, pulsante di contatto, badge disponibilità 24/7 e "
     "tutti i numeri di telefono cliccabili."),
    ("Chi Siamo", "Presentazione dell'agenzia con i valori fondanti: empatia, affidabilità "
     "e rispetto. Layout elegante con icone e testi professionali."),
    ("Servizi", "8 card interattive con icone dedicate per ogni servizio offerto: "
     "Disbrigo Pratiche, Vestizione Salme, Trasporti Nazionali ed Esteri, "
     "Camera Ardente, Addobbi Floreali, Manifesti Lutto, Servizi Cimiteriali, Cremazione."),
    ("Galleria", "6 fotografie professionali con griglia responsive e lightbox carousel "
     "integrato (navigazione con frecce, tastiera, swipe touch su mobile)."),
    ("Perché Sceglierci", "3 punti di forza: Disponibilità 24/7, Esperienza e Professionalità, "
     "Vicinanza e Rispetto."),
    ("Contatti", "Informazioni complete (indirizzo, 4 numeri di telefono cliccabili, email), "
     "modulo di contatto con validazione, mappa Google Maps interattiva."),
    ("Footer", "Riepilogo dati aziendali, link utili, tutti i recapiti, link social, "
     "privacy policy e cookie policy."),
]

for title, desc in sections_data:
    add_bullet(f": {desc}", title)

# 2.3
add_heading_styled("2.3 Responsive Design", level=2)
doc.add_paragraph(
    "Il sito si adatta perfettamente a qualsiasi dispositivo:"
)
add_bullet("Desktop: ", "layout completo con griglia a 3 colonne")
add_bullet("Tablet: ", "layout adattato a 2 colonne con margini ottimizzati")
add_bullet("Smartphone: ", "layout a colonna singola con menu hamburger dedicato")
add_bullet("Smartphone piccoli: ", "ulteriore ottimizzazione per schermi sotto i 480px")

# 2.4
add_heading_styled("2.4 Contatto Diretto WhatsApp", level=2)
doc.add_paragraph(
    "Pulsante WhatsApp sempre visibile in basso a destra, con messaggio precompilato. "
    "Permette ai clienti di contattarvi istantaneamente con un solo tocco, "
    "24 ore su 24, anche da dispositivi mobili."
)

# 2.5
add_heading_styled("2.5 Galleria Fotografica con Carousel", level=2)
doc.add_paragraph(
    "Galleria interattiva con 6 foto che si aprono in un lightbox a schermo intero. "
    "Il carousel integrato permette di navigare tra le immagini con:"
)
add_bullet("Frecce laterali (avanti/indietro con loop circolare)")
add_bullet("Frecce da tastiera (sinistra/destra)")
add_bullet("Swipe touch su dispositivi mobili")
add_bullet("Contatore foto (es. \"3 / 6\")")
add_bullet("Chiusura con click esterno, tasto X o Escape")

doc.add_page_break()

# ============================================================
# 3. TECNOLOGIA E PRESTAZIONI
# ============================================================
add_heading_styled("3. Tecnologia e Prestazioni", level=1)

doc.add_paragraph(
    "Il sito è stato sviluppato con tecnologie moderne e leggere, senza dipendenze "
    "da framework pesanti o CMS come WordPress:"
)

tech_data = [
    ("HTML5 Semantico", "Struttura accessibile e ottimizzata per i motori di ricerca"),
    ("CSS3 Moderno", "Variabili CSS, Flexbox, Grid Layout, animazioni native"),
    ("JavaScript Vanilla (ES6+)", "Nessuna libreria esterna, codice nativo ultra-leggero"),
    ("Google Fonts", "Tipografia professionale caricata in modo ottimizzato"),
    ("Font Awesome 6", "Libreria icone vettoriali di alta qualità"),
]

for title, desc in tech_data:
    add_bullet(f": {desc}", title)

add_heading_styled("Vantaggi tecnici rispetto a soluzioni WordPress/CMS:", level=3)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Caratteristica", "Sito Custom (Vostro)", "Sito WordPress"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

rows_data = [
    ["Velocità di caricamento", "< 1 secondo", "2-5 secondi"],
    ["Peso pagina", "~ 300 KB", "2-5 MB"],
    ["Sicurezza", "Nessuna vulnerabilità CMS", "Aggiornamenti continui richiesti"],
    ["Manutenzione", "Zero manutenzione tecnica", "Plugin/tema da aggiornare"],
    ["Costi hosting", "Gratuito (GitHub Pages)", "10-50 EUR/mese"],
]

for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

# ============================================================
# 4. CONFORMITA GDPR
# ============================================================
add_heading_styled("4. Conformità GDPR e Privacy", level=1)

doc.add_paragraph(
    "Il sito è pienamente conforme alla normativa europea sulla protezione dei dati "
    "personali (Regolamento UE 2016/679 - GDPR):"
)

add_bullet("Cookie Banner intelligente: ", "con 3 opzioni (Accetta tutti, Solo necessari, Personalizza)")
add_bullet("Pannello personalizzazione cookie: ", "con toggle per categorie (tecnici sempre attivi, terze parti opzionali)")
add_bullet("Google Maps condizionale: ", "caricato solo dopo consenso esplicito dell'utente")
add_bullet("Persistenza preferenze: ", "salvate nel browser per 6 mesi, poi richiedono nuovo consenso")
add_bullet("Privacy Policy completa: ", "pagina dedicata con 11 articoli conformi all'art. 13 del GDPR")
add_bullet("Cookie Policy dettagliata: ", "elenco cookie tecnici e terze parti, istruzioni disabilitazione browser")
add_bullet("Checkbox privacy nel form: ", "consenso obbligatorio prima dell'invio del messaggio")

# ============================================================
# 5. SEO
# ============================================================
add_heading_styled("5. Ottimizzazione SEO", level=1)

doc.add_paragraph(
    "Il sito è stato ottimizzato per il posizionamento sui motori di ricerca (Google):"
)

add_bullet("Meta tag ottimizzati: ", "title, description e keywords specifici per il settore")
add_bullet("Open Graph: ", "anteprima ottimizzata per la condivisione sui social media")
add_bullet("HTML semantico: ", "struttura heading corretta (H1, H2, H3) per la lettura da parte di Google")
add_bullet("Attributi alt su tutte le immagini: ", "per accessibilità e indicizzazione")
add_bullet("Velocità di caricamento: ", "fattore di ranking fondamentale per Google")
add_bullet("Design responsive: ", "Google privilegia i siti mobile-friendly nei risultati di ricerca")
add_bullet("Favicon e Apple Touch Icon: ", "generati dal logo per riconoscibilità nel browser")

doc.add_page_break()

# ============================================================
# 6. ANALISI COMPETITOR
# ============================================================
add_heading_styled("6. Analisi di Mercato e Competitor", level=1)

doc.add_paragraph(
    "Per garantire un posizionamento competitivo della proposta, è stata condotta un'analisi "
    "del mercato italiano dei servizi web per agenzie di onoranze funebri. "
    "Il mercato si suddivide in tre categorie principali:"
)

add_heading_styled("6.1 Piattaforme SaaS (Abbonamento ricorrente)", level=2)
doc.add_paragraph(
    "Soluzioni in abbonamento mensile/semestrale che combinano gestionale e sito web:"
)

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Piattaforma", "Costo", "Note"]):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

saas_data = [
    ["ModularSoftware", "250 EUR/anno", "Gestionale + sito template base"],
    ["Aurora365", "228-576 EUR/anno", "Solo gestionale, sito escluso"],
    ["Portale Funebre", "1.200 EUR/anno", "Gestionale + necrologi online"],
    ["OnoranzeFunebriCloud", "600-18.000 EUR/anno", "Soluzione modulare enterprise"],
]

for i, row_data in enumerate(saas_data):
    for j, val in enumerate(row_data):
        cell = table2.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

add_heading_styled("6.2 Web Agency (Sito custom una tantum)", level=2)

table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Light Shading Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Tipologia", "Range Prezzo", "Cosa include"]):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

agency_data = [
    ["Sito vetrina base", "800 - 1.400 EUR", "5-7 pagine WordPress, responsive"],
    ["Sito vetrina medio", "1.500 - 3.000 EUR", "Gallery, form, necrologi, SEO"],
    ["Sito completo", "3.000 - 5.000+ EUR", "E-commerce, area riservata, custom"],
]

for i, row_data in enumerate(agency_data):
    for j, val in enumerate(row_data):
        cell = table3.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph(
    "A questi costi vanno aggiunti i costi ricorrenti di hosting (10-50 EUR/mese), "
    "manutenzione annuale (300-2.000 EUR/anno) e dominio (10-30 EUR/anno)."
)

add_heading_styled("6.3 Il vostro vantaggio", level=2)
doc.add_paragraph(
    "Lavorando con un freelance specializzato, beneficiate di:"
)
add_bullet("Costo significativamente inferiore ", "rispetto a web agency strutturate")
add_bullet("Nessun intermediario: ", "comunicazione diretta con chi sviluppa il sito")
add_bullet("Codice proprietario: ", "nessun vincolo con piattaforme o abbonamenti")
add_bullet("Hosting gratuito: ", "GitHub Pages, zero costi di server")
add_bullet("Zero manutenzione tecnica: ", "sito statico, nessun plugin da aggiornare")
add_bullet("Qualità equivalente o superiore: ", "design custom, prestazioni imbattibili")

doc.add_page_break()

# ============================================================
# 7. PREVENTIVO
# ============================================================
add_heading_styled("7. Preventivo Economico", level=1)

doc.add_paragraph(
    "Di seguito il dettaglio dei costi per la realizzazione del sito web:"
)

# Tabella preventivo
table4 = doc.add_table(rows=12, cols=2)
table4.style = 'Light Shading Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

# Larghezza colonne
for cell in table4.columns[0].cells:
    cell.width = Cm(12)
for cell in table4.columns[1].cells:
    cell.width = Cm(4)

prev_headers = ["Voce", "Importo"]
for i, h in enumerate(prev_headers):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)

prev_data = [
    ["Progettazione grafica e UX/UI design", "150,00 EUR"],
    ["Sviluppo front-end (HTML5, CSS3, JavaScript)", "200,00 EUR"],
    ["Design responsive (desktop, tablet, mobile)", "80,00 EUR"],
    ["Animazioni e interazioni (particelle, scroll, carousel)", "50,00 EUR"],
    ["Galleria fotografica con lightbox carousel", "40,00 EUR"],
    ["Integrazione Google Maps e link di contatto", "30,00 EUR"],
    ["Conformità GDPR (cookie banner, privacy policy, cookie policy)", "60,00 EUR"],
    ["Ottimizzazione SEO base e meta tag", "40,00 EUR"],
    ["Generazione favicon e icone", "10,00 EUR"],
    ["Pubblicazione online e configurazione hosting", "40,00 EUR"],
]

for i, row_data in enumerate(prev_data):
    for j, val in enumerate(row_data):
        cell = table4.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Riga totale
total_row = table4.rows[11]
cell_label = total_row.cells[0]
cell_label.text = "TOTALE (IVA esclusa)"
for p in cell_label.paragraphs:
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

cell_total = total_row.cells[1]
cell_total.text = "700,00 EUR"
for p in cell_total.paragraphs:
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Box risparmio
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Risparmio rispetto alla media di mercato: oltre il 50%")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(Media web agency per sito equivalente: 1.500 - 3.000 EUR)")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_heading_styled("Costi ricorrenti", level=2)

table5 = doc.add_table(rows=4, cols=3)
table5.style = 'Light Shading Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Voce", "Costo", "Note"]):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

recurring_data = [
    ["Hosting (GitHub Pages)", "GRATUITO", "Incluso, nessun costo server"],
    ["Dominio .it (se richiesto)", "10-15 EUR/anno", "Opzionale, registrazione annuale"],
    ["Manutenzione/aggiornamenti", "Su richiesta", "Modifiche contenuti, nuove foto, ecc."],
]

for i, row_data in enumerate(recurring_data):
    for j, val in enumerate(row_data):
        cell = table5.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

# ============================================================
# 8. MODALITA DI PAGAMENTO
# ============================================================
add_heading_styled("8. Modalità di Pagamento", level=1)

add_bullet("50% all'accettazione del preventivo (acconto)")
add_bullet("50% alla consegna e pubblicazione del sito")
add_bullet("Pagamento tramite bonifico bancario o contanti")

# ============================================================
# 9. COSA INCLUDE
# ============================================================
add_heading_styled("9. Riepilogo: Cosa Include il Sito", level=1)

includes = [
    "Design esclusivo e personalizzato (no template)",
    "Sito single-page con 7 sezioni complete",
    "8 servizi dettagliati con icone",
    "Galleria fotografica con 6 foto e carousel lightbox",
    "Animazioni professionali (particelle, fade-in, hover)",
    "4 numeri di telefono cliccabili (click-to-call)",
    "Pulsante WhatsApp sempre visibile",
    "Modulo di contatto con validazione",
    "Mappa Google Maps interattiva",
    "Conformità GDPR completa (cookie banner + privacy + cookie policy)",
    "Design responsive per tutti i dispositivi",
    "Ottimizzazione SEO per Google",
    "Favicon personalizzato dal logo",
    "Hosting gratuito incluso",
    "Pubblicazione online e configurazione",
    "Codice sorgente di proprietà del cliente",
]

for item in includes:
    add_bullet(item)

# ============================================================
# 10. CONTATTI
# ============================================================
doc.add_page_break()
add_heading_styled("10. Contatti", level=1)

doc.add_paragraph(
    "Per qualsiasi chiarimento o per procedere con l'accettazione del preventivo, "
    "non esitate a contattarmi:"
)

doc.add_paragraph()
add_bold_paragraph("Enrico Maria Caruso")
add_bold_paragraph("Full Stack Developer & Web Designer Freelance")
doc.add_paragraph()
add_bold_paragraph("Sito web: ", "www.emcdigitalsolutions.it")
doc.add_paragraph()

p = doc.add_paragraph()
p.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Grazie per la fiducia.")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)
run.italic = True

# ============================================================
# SALVATAGGIO
# ============================================================
output_path = r"C:\workspace\giulianalagreca\Preventivo_Sito_Web_Giuliana_La_Greca.docx"
doc.save(output_path)
print(f"Documento salvato: {output_path}")
